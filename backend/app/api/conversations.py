import logging
from uuid import UUID
from fastapi import APIRouter, Depends, HTTPException, Query
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, desc, update
from app.database import get_db
from app.models import Conversation, Message, User, SentEmail, ConvMemory
from app.api.deps import get_current_user
from app.schemas.conversation import ConversationItem, MessageItem, UpdateTitleRequest
from app.services.llm_service import llm_router
from app.services.email_service import send_email
from app.domain.conversation.service import ConversationService

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/conversations", tags=["conversations"])


@router.get("")
async def list_conversations(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, le=100),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> dict:
    offset = (page - 1) * page_size
    total_result = await db.execute(
        select(func.count(Conversation.id)).where(Conversation.user_id == current_user.id)
    )
    total = total_result.scalar()
    result = await db.execute(
        select(Conversation)
        .where(Conversation.user_id == current_user.id)
        .order_by(desc(Conversation.updated_at))
        .offset(offset).limit(page_size)
    )
    convs = result.scalars().all()

    # Fetch last message for each conversation
    conv_ids = [c.id for c in convs]
    if conv_ids:
        # Subquery: get newest message per conversation
        subq = (
            select(
                Message.conv_id,
                Message.content,
                func.row_number().over(
                    partition_by=Message.conv_id,
                    order_by=desc(Message.created_at),
                ).label("rn")
            )
            .where(Message.conv_id.in_(conv_ids))
            .subquery()
        )
        last_msgs = await db.execute(
            select(subq.c.conv_id, subq.c.content).where(subq.c.rn == 1)
        )
        last_msg_map = {row.conv_id: row.content for row in last_msgs}
    else:
        last_msg_map = {}

    return {
        "items": [
            ConversationItem(
                id=str(c.id),
                title=c.title,
                last_message=last_msg_map.get(c.id),
                created_at=c.created_at,
                updated_at=c.updated_at,
            )
            for c in convs
        ],
        "total": total,
        "page": page,
        "page_size": page_size,
    }


@router.get("/{conv_id}/messages")
async def list_messages(
    conv_id: UUID,
    cursor: str = Query(None),
    limit: int = Query(20, le=100),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> dict:
    conv_result = await db.execute(
        select(Conversation).where(Conversation.id == conv_id, Conversation.user_id == current_user.id)
    )
    if not conv_result.scalar_one_or_none():
        raise HTTPException(404, "Conversation not found")

    q = select(Message).where(Message.conv_id == conv_id).order_by(desc(Message.created_at))
    if cursor:
        q = q.where(Message.created_at < cursor)
    q = q.limit(limit)
    result = await db.execute(q)
    msgs = result.scalars().all()
    return {
        "items": [
            MessageItem(id=str(m.id), role=m.role.value, type=m.type.value,
                       content=m.content, audio_url=m.audio_url, created_at=m.created_at)
            for m in reversed(msgs)
        ],
        "next_cursor": str(msgs[-1].created_at) if len(msgs) == limit else None,
    }


@router.get("/{conv_id}/memory")
async def get_conv_memory(
    conv_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> dict:
    """获取对话级记忆摘要"""
    conv_result = await db.execute(
        select(Conversation).where(
            Conversation.id == conv_id,
            Conversation.user_id == current_user.id,
        )
    )
    if not conv_result.scalar_one_or_none():
        raise HTTPException(404, "Conversation not found")

    result = await db.execute(
        select(ConvMemory.summary_text)
        .where(ConvMemory.conv_id == conv_id)
        .order_by(ConvMemory.updated_at.desc())
        .limit(1)
    )
    row = result.scalar_one_or_none()
    return {"summary_text": row or ""}


@router.get("/summaries-all")
async def list_summaries(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> dict:
    """List all conversation summaries (conv_memories) for the current user."""
    result = await db.execute(
        select(ConvMemory, Conversation.title)
        .join(Conversation, ConvMemory.conv_id == Conversation.id)
        .where(ConvMemory.user_id == current_user.id)
        .where(ConvMemory.summary_text != "")
        .order_by(desc(ConvMemory.updated_at))
        .limit(50)
    )
    rows = result.all()
    return {
        "items": [
            {
                "conv_id": str(row.ConvMemory.conv_id),
                "conv_title": row.title,
                "summary_text": row.ConvMemory.summary_text,
                "updated_at": row.ConvMemory.updated_at.isoformat(),
            }
            for row in rows
        ]
    }


@router.delete("/{conv_id}")
async def delete_conversation(
    conv_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Conversation).where(
            Conversation.id == conv_id,
            Conversation.user_id == current_user.id,
        )
    )
    conv = result.scalar_one_or_none()
    if not conv:
        raise HTTPException(404, "Conversation not found")
    # Nullify source_conv_id in related user_memories to avoid FK violation
    from app.models.user_memory import UserMemory
    await db.execute(
        update(UserMemory)
        .where(UserMemory.source_conv_id == conv_id)
        .values(source_conv_id=None)
    )
    await db.delete(conv)
    await db.commit()
    return {"status": "deleted"}


@router.put("/{conv_id}/title")
async def update_title(
    conv_id: UUID,
    req: UpdateTitleRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Conversation).where(Conversation.id == conv_id, Conversation.user_id == current_user.id)
    )
    conv = result.scalar_one_or_none()
    if not conv:
        raise HTTPException(404, "Conversation not found")
    conv.title = req.title
    await db.commit()
    return {"status": "ok"}


@router.post("/{conv_id}/summary")
async def get_conversation_summary(
    conv_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Summarize the conversation via LLM and return the summary text."""
    conv_result = await db.execute(
        select(Conversation).where(
            Conversation.id == conv_id,
            Conversation.user_id == current_user.id,
        )
    )
    if not conv_result.scalar_one_or_none():
        raise HTTPException(404, "Conversation not found")

    try:
        summary = await ConversationService.generate_summary(conv_id, db, llm_router)
        return {"summary": summary}
    except ValueError as e:
        raise HTTPException(400, str(e))
    except RuntimeError as e:
        raise HTTPException(500, str(e))


@router.post("/{conv_id}/email-summary")
async def email_summary(
    conv_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Summarize the conversation via LLM and send it to the user's email."""
    conv_result = await db.execute(
        select(Conversation).where(
            Conversation.id == conv_id,
            Conversation.user_id == current_user.id,
        )
    )
    conv = conv_result.scalar_one_or_none()
    if not conv:
        raise HTTPException(404, "Conversation not found")

    try:
        email = await ConversationService.email_summary(
            conv_id, current_user.id, current_user.email or "",
            conv.title, db, llm_router, send_email,
        )
        return {"status": "sent", "email": email}
    except ValueError as e:
        raise HTTPException(400, str(e))
    except RuntimeError as e:
        raise HTTPException(500, str(e))


@router.post("/{conv_id}/diary")
async def generate_diary(
    conv_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Generate an AI diary entry from the day's conversation."""
    conv_result = await db.execute(
        select(Conversation).where(
            Conversation.id == conv_id,
            Conversation.user_id == current_user.id,
        )
    )
    conv = conv_result.scalar_one_or_none()
    if not conv:
        raise HTTPException(404, "Conversation not found")

    msgs_result = await db.execute(
        select(Message)
        .where(Message.conv_id == conv_id)
        .order_by(Message.created_at)
    )
    messages = msgs_result.scalars().all()
    if len(messages) < 2:
        raise HTTPException(400, "对话内容太少，多聊几句吧")

    dialogue = "\n".join(
        f"{'🧑' if m.role.value == 'user' else '🐱'}: {m.content or ''}"
        for m in messages
    )

    prompt = (
        f"把以下对话写成一段温暖的日记（100-150字，第一人称）。"
        f"像写给自己看的，自然不做作。包含：今天聊了什么、心情如何、有什么收获。\n\n"
        f"对话:\n{dialogue}\n\n日记:"
    )
    try:
        diary = await llm_router.chat([{"role": "user", "content": prompt}])
        diary = (diary or "").strip()
    except Exception:
        raise HTTPException(500, "生成失败")

    from app.models.note import Note
    note = Note(
        user_id=current_user.id,
        title=f"📔 {conv.title}",
        content=diary,
        note_type="diary",
    )
    db.add(note)
    await db.commit()
    await db.refresh(note)

    return {"id": str(note.id), "title": note.title, "content": diary}


@router.get("/sent-emails")
async def list_sent_emails(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """List all sent email summaries for the current user."""
    result = await db.execute(
        select(SentEmail)
        .where(SentEmail.user_id == current_user.id)
        .order_by(SentEmail.sent_at.desc())
        .limit(20)
    )
    records = result.scalars().all()
    return {
        "items": [
            {
                "id": str(r.id),
                "conv_title": r.conv_title,
                "recipient": r.recipient,
                "summary_preview": r.summary_preview,
                "sent_at": r.sent_at.isoformat(),
            }
            for r in records
        ]
    }


@router.post("/{conv_id}/export")
async def export_conversation(
    conv_id: UUID,
    format: str = Query("pdf", regex="^(pdf|docx)$"),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Export conversation as PDF or DOCX file."""
    conv_result = await db.execute(
        select(Conversation).where(
            Conversation.id == conv_id,
            Conversation.user_id == current_user.id,
        )
    )
    conv = conv_result.scalar_one_or_none()
    if not conv:
        raise HTTPException(404, "Conversation not found")

    msgs_result = await db.execute(
        select(Message)
        .where(Message.conv_id == conv_id)
        .order_by(Message.created_at)
    )
    messages = msgs_result.scalars().all()
    if not messages:
        raise HTTPException(400, "对话内容为空")

    lines = []
    for m in messages:
        role = "用户" if m.role.value == "user" else "灵犀"
        lines.append(f"{role}: {m.content or ''}")
    dialogue = "\n\n".join(lines)
    title = conv.title or "对话"

    import io as io_mod, uuid as uuid_mod
    from app.services.minio_service import upload_file, get_download_url

    file_bytes = io_mod.BytesIO()
    filename = f"{title}.{format}"

    if format == "pdf":
        from fpdf import FPDF
        from app.services.conversion_service import _find_cjk_font
        pdf = FPDF()
        pdf.add_page()
        font_path = _find_cjk_font()
        if font_path:
            pdf.add_font("cjk", "", font_path, uni=True)
            font_name = "cjk"
        else:
            font_name = "Courier"
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font(font_name, "", 12)
        pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(6)
        for line in dialogue.split("\n"):
            if not line.strip():
                pdf.ln(3)
                continue
            pdf.set_font(font_name, "", 9)
            pdf.multi_cell(0, 5.5, line)
        pdf.output(file_bytes)
    else:
        from docx import Document
        doc = Document()
        doc.add_heading(title, 0)
        for line in dialogue.split("\n"):
            doc.add_paragraph(line)
        doc.save(file_bytes)

    file_bytes.seek(0)
    file_size = len(file_bytes.getvalue())

    object_name = await upload_file(
        file_bytes.getvalue(), filename,
        content_type="application/pdf" if format == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    if not object_name:
        raise HTTPException(500, "文件上传失败")

    download_url_val = await get_download_url(object_name)
    if not download_url_val:
        raise HTTPException(500, "文件上传失败")

    from app.models import ConvertedFile
    target_name = f"{title}.{format}"
    cf = ConvertedFile(
        user_id=current_user.id,
        original_name=f"{title}.txt",
        target_name=target_name,
        object_name=object_name,
        file_size=file_size,
        content_type="application/pdf" if format == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    db.add(cf)
    await db.commit()
    await db.refresh(cf)

    return {
        "id": str(cf.id),
        "target_name": target_name,
        "download_url": download_url_val,
    }
